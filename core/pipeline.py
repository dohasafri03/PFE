"""
Pipeline v2 - Lightweight orchestrator
Scrape CSV -> Filter/Score -> Analyse NLP -> Generate Dossiers

Works directly with CSV data, no database required.
Dependencies: jinja2, python-docx, spacy (fr_core_news_md)
"""
import csv
import json
import logging
import os
import re
import sys
import uuid
from dataclasses import dataclass, field, asdict
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple

os.environ['PYTHONIOENCODING'] = 'utf-8'

logger = logging.getLogger(__name__)


# ============================================================
# Data models (lightweight, no DB)
# ============================================================

class Priority(Enum):
    HOT = "HOT"
    WARM = "WARM"
    COLD = "COLD"
    EXCLUDED = "EXCLUDED"


class Domain(Enum):
    DEV = "dev"
    DATA = "data"
    AI = "ai"
    CLOUD = "cloud"


@dataclass
class Consultation:
    """A single consultation/tender from the CSV."""
    id: str = ""
    reference: str = ""
    objet: str = ""
    acheteur: str = ""
    categorie: str = ""
    nature_prestation: str = ""
    date_publication: str = ""
    date_limite: str = ""
    lieu_execution: str = ""
    budget_estime: str = ""
    nb_articles: str = ""
    articles: str = ""
    fichiers_joints: str = ""
    email_contact: str = ""
    url: str = ""
    source: str = ""
    # Scoring fields (filled by pipeline)
    priority: Priority = Priority.EXCLUDED
    score_total: int = 0
    domain_scores: Dict[str, int] = field(default_factory=dict)
    matched_keywords: List[str] = field(default_factory=list)
    matched_domains: List[str] = field(default_factory=list)
    is_excluded: bool = False
    # NLP analysis (filled by pipeline)
    nlp_categories: List[str] = field(default_factory=list)
    nlp_confidence: float = 0.0
    nlp_key_entities: List[str] = field(default_factory=list)
    nlp_summary: str = ""
    # Generated descriptions (filled by pipeline)
    description_technique: str = ""
    description_fonctionnelle: str = ""
    requirements: List[str] = field(default_factory=list)
    # Enrichment from CPS/detail pages (filled by pipeline stage 2.5)
    estimation_budget: str = ""
    caution_provisoire: str = ""
    procedure: str = ""
    allotissement: str = ""
    domaines_activite: str = ""
    qualifications: str = ""
    reservation_pme: str = ""
    contenu_cps: str = ""  # Full extracted text from CPS (HTML or PDF/OCR)
    cps_source: str = ""  # 'prado_html', 'bdc_pdf', 'bdc_ocr', ''
    # Recommendations (filled by pipeline)
    strengths: List[str] = field(default_factory=list)
    risks: List[str] = field(default_factory=list)
    recommendations: List[str] = field(default_factory=list)

    @classmethod
    def from_csv_row(cls, row: dict) -> "Consultation":
        """Create from a CSV DictReader row."""
        return cls(
            id=row.get('id', ''),
            reference=row.get('reference', ''),
            objet=row.get('objet', ''),
            acheteur=row.get('acheteur', ''),
            categorie=row.get('categorie', ''),
            nature_prestation=row.get('nature_prestation', ''),
            date_publication=row.get('date_publication', ''),
            date_limite=row.get('date_limite', ''),
            lieu_execution=row.get('lieu_execution', ''),
            budget_estime=row.get('budget_estime', ''),
            nb_articles=row.get('nb_articles', ''),
            articles=row.get('articles', ''),
            fichiers_joints=row.get('fichiers_joints', ''),
            email_contact=row.get('email_contact', ''),
            url=row.get('url', ''),
            source=row.get('source', ''),
        )

    def searchable_text(self) -> str:
        """Concatenate all text fields for scoring/NLP."""
        return ' '.join([
            self.objet, self.articles, self.nature_prestation,
            self.reference, self.acheteur
        ]).lower()


# ============================================================
# Weighted scoring engine (from filter_alexsys_v3 logic)
# ============================================================

ALEXSYS_KEYWORDS = {
    'dev': [
        (r'\blogiciel\b', 3), (r'\bprogiciel\b', 3), (r'\berp\b', 3), (r'\bcrm\b', 3),
        (r'\bsyst[eè]me\s+d.information\b', 3),
        (r'\bd[eé]veloppement\s+(?:web|mobile|logiciel|informatique|applicat)', 3),
        (r'\bapplication\s+(?:mobile|web)\b', 3),
        (r'\brefonte\s+(?:du\s+)?site\s+web\b', 3),
        (r'\bsite\s+web\b', 2),
        (r'\bplateforme\s+(?:num[eé]rique|digitale|web|informatique)\b', 2),
        (r'\bportail\s+(?:web|num[eé]rique|digital)\b', 2),
        (r'\bint[eé]gration\s+(?:syst[eè]me|logiciel|informatique)\b', 2),
        (r'\bmigration\s+(?:syst[eè]me|logiciel|donn[eé]es|informatique)\b', 2),
        (r'\bmise\s+en\s+(?:oeuvre|place|service).*(?:logiciel|solution|syst[eè]me|plateforme)', 2),
        (r'\bmaintenance\s+(?:informatique|logiciel|syst[eè]me)', 2),
        (r'\bentretien\s+(?:logiciel|informatique)', 2),
        (r'\btransformation\s+(?:num[eé]rique|digitale)\b', 2),
        (r'\bdigitalisation\b', 2),
        (r'\binfrastructure\s+(?:it|informatique|num[eé]rique)\b', 2),
        (r'\bcybers[eé]curit[eé]\b', 3),
        (r'\bs[eé]curit[eé]\s+informatique\b', 2),
        (r'\bsolution\s+de\s+s[eé]curit[eé]\b', 2),
        (r'\bfirewall\b', 2), (r'\bantivirus\b', 2),
        (r'\binfog[eé]rance\b', 3),
        (r'\baudit\s+(?:informatique|syst[eè]me|s[eé]curit[eé])', 2),
        (r'\bh[eé]bergement\s+(?:web|informatique|cloud|serveur)', 2),
    ],
    'data': [
        (r'\bbase\s+de\s+donn[eé]es\b', 2),
        (r'\bbusiness\s+intelligence\b', 3), (r'\banalytics\b', 3),
        (r'\breporting\b', 2), (r'\bdashboard\b', 3),
        (r'\btableau\s+de\s+bord\b', 2), (r'\bbig\s*data\b', 3),
        (r'\bdatawarehouse\b', 3), (r'\betl\b', 3),
        (r'\bgestion\s+(?:des\s+)?donn[eé]es\b', 2),
        (r'\btraitement\s+(?:de|des)\s+donn[eé]es\b', 2),
        (r'\bdata\s+(?:center|management|lake|warehouse)\b', 2),
        (r'\bs[eé]curisation\s+(?:des\s+)?donn[eé]es\b', 2),
    ],
    'ai': [
        (r'\bintelligence\s+artificielle\b', 3),
        (r'\bmachine\s+learning\b', 3), (r'\bdeep\s+learning\b', 3),
        (r'\balgorithme\b', 2), (r'\bpr[eé]dictif\b', 2),
        (r'\bautomatisation\b', 1), (r'\brobot(?:ique|isation)\b', 2),
        (r'\bchatbot\b', 3), (r'\bnlp\b', 3),
    ],
    'cloud': [
        (r'\bcloud\b', 3), (r'\bsaas\b', 3), (r'\bpaas\b', 3), (r'\biaas\b', 3),
        (r'\bvirtualisation\b', 3), (r'\bdatacenter\b', 2), (r'\bdata\s*center\b', 2),
        (r'\binfrastructure\s+cloud\b', 3), (r'\bdevops\b', 3),
        (r'\bserveur\b', 1), (r'\bsauvegarde\b', 1),
    ],
}

EXCLUSION_PATTERNS = [
    r"\bavis d'?annulation\b",
    r"\bannul[ée]\b",
    r'\b[eé]lectrification\b',
    r'\br[eé]seau\s+(?:bt|basse|haute|mt\b|eau|assain)',
    r'\btravaux\s+(?:de\s+)?(?:construction|b[aâ]timent|route|voirie|peinture)',
    r'\bbranchement\s+(?:eau|[eé]lectr)',
    r'\bmat[eé]riaux\s+(?:de\s+)?construction\b',
    r'\barticle\s+(?:de\s+)?sport\b',
    r'\bproduit\s+(?:alimentaire|d.entretien|pharmaceutique|chimique)',
    r'\bmeuble\b', r'\bmobilier\b', r'\bv[eé]hicule\b',
    r'\blocation\s+(?:de\s+)?(?:voiture|v[eé]hicule|moyen)',
    r'\b(?:engrais|semence|b[eé]tail|fourrage|irrigation)\b',
    r'\b(?:m[eé]dicament|vaccin|r[eé]actif|proth[eè]se)\b',
    r'\bproduit\s+d.entretien\b',
]

OVERRIDE_PATTERNS = [
    r'\blogiciel\b', r'\binformatique\b', r'\berp\b', r'\bcrm\b',
    r'\bcybers[eé]curit[eé]\b', r'\bvirtualisation\b', r'\bcloud\b',
    r'\bsite\s+web\b', r'\bapplication\b', r'\bplateforme\b',
    r'\bsyst[eè]me\s+d.information\b', r'\binfog[eé]rance\b',
]


class ScoringEngine:
    """Weighted keyword scoring for IT relevance."""

    @staticmethod
    def is_excluded(text: str) -> bool:
        # Hard-exclude cancelled notices, even if they contain IT keywords.
        if re.search(r"\bavis\s+d[’']?annulation\b", text, re.I) or re.search(r"\bdate\s+d[’']?annulation\b", text, re.I):
            return True
        for pat in EXCLUSION_PATTERNS:
            if re.search(pat, text, re.I):
                for ovr in OVERRIDE_PATTERNS:
                    if re.search(ovr, text, re.I):
                        return False
                return True
        return False

    @staticmethod
    def score(consultation: Consultation) -> None:
        """Score a consultation in-place."""
        text = consultation.searchable_text()

        if ScoringEngine.is_excluded(text):
            consultation.is_excluded = True
            consultation.priority = Priority.EXCLUDED
            return

        domain_scores = {}
        total = 0
        matched = []

        for domain, patterns in ALEXSYS_KEYWORDS.items():
            dw = 0
            for pattern, weight in patterns:
                if re.search(pattern, text, re.I):
                    dw += weight
                    # Clean pattern for display
                    kw = re.sub(r'\\b|\\s\+|\(\?[^)]*\)|\\s|\[\^.*?\]', '', pattern).strip()
                    matched.append(f"{domain}:{kw}(w{weight})")
            domain_scores[domain] = dw
            total += dw

        consultation.score_total = total
        consultation.domain_scores = domain_scores
        consultation.matched_keywords = matched
        consultation.matched_domains = [d for d, w in domain_scores.items() if w > 0]

        if total >= 4:
            consultation.priority = Priority.HOT
        elif total >= 2:
            consultation.priority = Priority.WARM
        elif total >= 1:
            consultation.priority = Priority.COLD
        else:
            consultation.priority = Priority.EXCLUDED


# ============================================================
# CPS Extractor: download & extract CPS content
# ============================================================

class CPSExtractor:
    """Extract detailed info from consultation detail pages.

    - PRADO pages: scrape structured HTML (no auth needed for metadata)
    - BDC pages: download ZIP, extract PDF, OCR if scanned
    """

    BASE = "https://www.marchespublics.gov.ma"

    def __init__(self):
        import requests as _req
        self.session = _req.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Accept-Language': 'fr-FR,fr;q=0.9,en;q=0.8',
        })
        # Warm up session (get cookies)
        try:
            self.session.get(self.BASE, timeout=15)
        except Exception:
            pass

    def enrich(self, consultation: Consultation) -> bool:
        """Enrich a consultation with data from its detail page. Returns True if enriched."""
        url = consultation.url
        if not url:
            return False
        try:
            if '/bdc/' in url:
                return self._enrich_bdc(consultation)
            else:
                return self._enrich_prado(consultation)
        except Exception as e:
            logger.debug(f"CPSExtractor error for {consultation.reference}: {e}")
            return False

    # ---- PRADO portal (old portal, structured HTML) ----
    def _enrich_prado(self, c: Consultation) -> bool:
        """Extract structured data from PRADO consultation detail page."""
        from bs4 import BeautifulSoup
        r = self.session.get(c.url, timeout=20)
        if r.status_code != 200:
            return False

        soup = BeautifulSoup(r.text, 'html.parser')
        text = soup.get_text(separator='\n', strip=True)

        # Extract key-value fields from the page text
        # NOTE: PRADO pages render labels and values on separate lines.
        # Some fields (like Estimation) may have 1-2 noise lines (* / :) between
        # the label and the actual numeric value.
        field_patterns = {
            'estimation_budget': [
                # Label then skip non-digit lines (*, :, etc.) then capture number
                r'Estimation\s*\(en Dhs TTC\)\s*\n(?:[^\d\n]*\n)*\s*([\d][\d\s,.]+)',
                r'Estimation[^:\n]*?:\s*\n(?:[^\d\n]*\n)*\s*([\d][\d\s,.]+)',
                r'Estimation[^:\n]*?:\s*([\d\s,.]+(?:MAD|DH|Dhs[^\n]{0,20}))',
            ],
            'caution_provisoire': [
                r'Caution provisoire\s*:\s*\n\s*([\d\s,.]+\s*(?:MAD|DH|Dhs)?[^\n]*)',
                r'Caution provisoire\s*:\s*([\d\s,.]+\s*(?:MAD|DH|Dhs)?)',
            ],
            'procedure': [
                r'Proc.dure\s*:\s*\n\s*([^\n]+)',
                r'Proc.dure\s*:\s*([^\n]+)',
            ],
            'allotissement': [
                r'Allotissement\s*:\s*\n\s*([^\n]+)',
                r'Allotissement\s*:\s*([^\n]+)',
            ],
            'domaines_activite': [
                r"Domaines?\s*d.activit.\s*:\s*\n\s*([^\n]+(?:\n[^\n]{5,100})?)",
                r"Domaines?\s*d.activit.\s*:\s*([^\n]+(?:\n[^\n]{5,60})?)",
            ],
            'qualifications': [
                r'Qualifications?\s*:\s*\n\s*([^\n]+(?:\n[^\n]{5,80}){0,3})',
                r'Qualifications?\s*:\s*([^\n]+(?:\n[^\n]{5,80}){0,3})',
            ],
            'reservation_pme': [
                r'(R.serv.\s+.+?(?:PME|TPE|jeunes|coop.ratives|auto-entrepreneurs)[^\n]*)',
            ],
        }

        for field_name, patterns in field_patterns.items():
            for pat in patterns:
                m = re.search(pat, text, re.I)
                if m:
                    val = m.group(1).strip()
                    # Clean trailing noise (next field labels that got captured)
                    val = re.split(
                        r'\n(?:Domaines?\s*d.activit|Agr.ments?\s*:|Qualifications?\s*:|'
                        r'Prospectus|Date et heure|R.union\s*:|Visite|'
                        r'Adresse de|Lieu d.ex|Cat.gorie|Lieu d.ouverture|'
                        r'Prix d.acquisition|Contact)',
                        val
                    )[0].strip()
                    # Skip empty/placeholder values
                    if val and len(val) > 1 and val not in ('-', ':'):
                        setattr(c, field_name, val)
                        break

        # Extract full page content as CPS context (remove navigation/menu)
        # Keep only the main content area
        main_part = soup.find('div', {'class': 'content'}) or soup.find('div', id='main-part') or soup.find('div', id='content')
        if main_part:
            for tag in main_part.find_all(['script', 'style', 'nav']):
                tag.decompose()
            content = main_part.get_text(separator='\n', strip=True)
        else:
            content = text

        # Clean: remove short lines, menu items, navigation chrome
        noise_keywords = [
            'javascript', 'aller au', 'cookie', "s'identifier",
            'mon panier', 'recherche avanc', 'nouvelle recherche',
            'consultations de test', "guides d'utilisation",
            "conditions d'utilisation", 'inforgaa', 'inforgs',
            'nos_partenaires', 'infosite', 'recherche rapide',
            'bon de commande', 'toutes les consultations',
            'toutes les annonces', 'langue de navigation',
            'vous n\'etes pas', 'vous n\'êtes pas',
            "sans retrait", "sans question", "sans message",
            "sans dépôt", "consultations clôturées",
            "fourniture d'écrans", # search example text
            'ex 1:', 'ex 2 :',
            'annonce de programme', 'annonce de synth',
            'liste des bons', 'liste des march',
            'liste des conventions',
        ]
        lines = []
        for line in content.split('\n'):
            line = line.strip()
            if len(line) > 15 and not any(kw in line.lower() for kw in noise_keywords):
                lines.append(line)

        c.contenu_cps = '\n'.join(lines[:200])  # Cap at 200 lines
        c.cps_source = 'prado_html'
        return True

    # ---- BDC portal (download ZIP -> PDF -> text/OCR) ----
    def _enrich_bdc(self, c: Consultation) -> bool:
        """Download files from BDC page and extract text."""
        from bs4 import BeautifulSoup
        import zipfile
        import io

        r = self.session.get(c.url, timeout=20)
        if r.status_code != 200:
            return False

        soup = BeautifulSoup(r.text, 'html.parser')

        # Find download links
        dl_links = soup.find_all('a', href=re.compile(r'/download/'))
        if not dl_links:
            # BDC page may also need login
            if 'connecter' in soup.get_text().lower():
                c.cps_source = 'bdc_auth_required'
            return False

        # Download first available file
        for a in dl_links:
            href = a['href']
            if not href.startswith('http'):
                href = self.BASE + href

            try:
                dr = self.session.get(href, timeout=60)
                if dr.status_code != 200:
                    continue
                ct = dr.headers.get('Content-Type', '')

                extracted_text = ''

                if 'zip' in ct or dr.content[:4] == b'PK\x03\x04':
                    extracted_text = self._extract_from_zip(dr.content)
                elif 'pdf' in ct or dr.content[:4] == b'%PDF':
                    extracted_text = self._extract_pdf_text(dr.content)
                elif 'word' in ct or 'docx' in ct:
                    extracted_text = self._extract_docx_text(dr.content)

                if extracted_text.strip():
                    c.contenu_cps = extracted_text[:15000]  # Cap at 15K chars
                    c.cps_source = 'bdc_pdf'
                    return True
                else:
                    # Try OCR
                    ocr_text = self._ocr_from_content(dr.content, ct)
                    if ocr_text.strip():
                        c.contenu_cps = ocr_text[:15000]
                        c.cps_source = 'bdc_ocr'
                        return True

            except Exception as e:
                logger.debug(f"BDC download error: {e}")
                continue

        return False

    def _extract_from_zip(self, zip_bytes: bytes) -> str:
        """Extract text from files inside a ZIP archive."""
        import zipfile
        import io
        texts = []
        try:
            z = zipfile.ZipFile(io.BytesIO(zip_bytes))
            for name in z.namelist():
                if name.startswith('__MACOSX'):
                    continue
                ext = name.lower().rsplit('.', 1)[-1] if '.' in name else ''
                data = z.read(name)
                if ext == 'pdf':
                    t = self._extract_pdf_text(data)
                    if t.strip():
                        texts.append(t)
                    else:
                        # Scanned PDF -> try OCR
                        ocr = self._ocr_pdf(data)
                        if ocr:
                            texts.append(ocr)
                elif ext == 'docx':
                    texts.append(self._extract_docx_text(data))
                elif ext == 'txt':
                    try:
                        texts.append(data.decode('utf-8-sig'))
                    except Exception:
                        texts.append(data.decode('latin-1', errors='replace'))
        except Exception as e:
            logger.debug(f"ZIP extraction error: {e}")
        return '\n\n'.join(texts)

    def _extract_pdf_text(self, pdf_bytes: bytes) -> str:
        """Extract text from a PDF using PyMuPDF."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=pdf_bytes, filetype='pdf')
            text = ''
            for page in doc:
                text += page.get_text()
            doc.close()
            return text.strip()
        except Exception as e:
            logger.debug(f"PDF extraction error: {e}")
            return ''

    def _extract_docx_text(self, docx_bytes: bytes) -> str:
        """Extract text from a DOCX file."""
        try:
            from docx import Document
            import io
            doc = Document(io.BytesIO(docx_bytes))
            return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            logger.debug(f"DOCX extraction error: {e}")
            return ''

    def _ocr_pdf(self, pdf_bytes: bytes) -> str:
        """OCR a scanned PDF using Tesseract."""
        try:
            import fitz
            doc = fitz.open(stream=pdf_bytes, filetype='pdf')
            images = []
            for page in doc:
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes('png')
                images.append(img_bytes)
            doc.close()
            return self._ocr_images(images)
        except Exception as e:
            logger.debug(f"OCR PDF error: {e}")
            return ''

    def _ocr_from_content(self, content: bytes, content_type: str) -> str:
        """Try OCR on raw content (ZIP or PDF)."""
        if b'PK\x03\x04' == content[:4] or 'zip' in content_type:
            import zipfile, io
            try:
                z = zipfile.ZipFile(io.BytesIO(content))
                for name in z.namelist():
                    if name.lower().endswith('.pdf') and not name.startswith('__'):
                        return self._ocr_pdf(z.read(name))
            except Exception:
                pass
        elif content[:4] == b'%PDF' or 'pdf' in content_type:
            return self._ocr_pdf(content)
        return ''

    def _ocr_images(self, image_bytes_list: list) -> str:
        """Run Tesseract OCR on a list of PNG image bytes."""
        try:
            from PIL import Image
            import pytesseract
            import io
            texts = []
            for img_bytes in image_bytes_list:
                img = Image.open(io.BytesIO(img_bytes))
                text = pytesseract.image_to_string(img, lang='fra+ara')
                if text.strip():
                    texts.append(text.strip())
            return '\n\n'.join(texts)
        except ImportError:
            logger.warning("pytesseract/Pillow not installed - OCR disabled")
            return ''
        except Exception as e:
            logger.debug(f"OCR error: {e}")
            return ''


# ============================================================
# Lightweight NLP analysis (spaCy)
# ============================================================

class NLPAnalyzer:
    """Lightweight NLP using spaCy for entity extraction and text analysis."""

    def __init__(self):
        self._nlp = None

    @property
    def nlp(self):
        if self._nlp is None:
            import spacy
            self._nlp = spacy.load("fr_core_news_md")
        return self._nlp

    def analyze(self, consultation: Consultation) -> None:
        """Run NLP analysis on a consultation, update in-place."""
        text = f"{consultation.objet}. {consultation.nature_prestation}. {consultation.articles}"
        text = text.strip()
        if not text or text == '. .':
            return

        doc = self.nlp(text[:5000])  # Limit to 5000 chars

        # Extract named entities
        entities = []
        for ent in doc.ents:
            if ent.label_ in ('ORG', 'PRODUCT', 'LOC', 'MISC') and len(ent.text) > 2:
                entities.append(f"{ent.text} ({ent.label_})")
        consultation.nlp_key_entities = list(set(entities))[:10]

        # Auto-categorize based on matched domains
        cat_map = {
            'dev': 'Developpement/Logiciel',
            'data': 'Data/BI',
            'ai': 'IA/Machine Learning',
            'cloud': 'Cloud/Infrastructure'
        }
        consultation.nlp_categories = [cat_map[d] for d in consultation.matched_domains if d in cat_map]

        # Simple confidence based on score
        if consultation.score_total >= 6:
            consultation.nlp_confidence = 0.95
        elif consultation.score_total >= 4:
            consultation.nlp_confidence = 0.85
        elif consultation.score_total >= 2:
            consultation.nlp_confidence = 0.70
        else:
            consultation.nlp_confidence = 0.50

        # Generate brief summary
        sentences = [s.text.strip() for s in doc.sents if len(s.text.strip()) > 10]
        consultation.nlp_summary = ' '.join(sentences[:3])[:300]

    def generate_descriptions(self, consultation: Consultation) -> None:
        """Generate description_technique, description_fonctionnelle, requirements."""
        objet = consultation.objet
        nature = consultation.nature_prestation
        articles = consultation.articles
        domains = consultation.matched_domains
        cat = consultation.categorie

        # --- Description technique ---
        tech_parts = []
        tech_parts.append(f"Le present marche porte sur: {objet}.")
        if nature:
            tech_parts.append(f"Nature de la prestation: {nature}.")

        # CPS-enriched info
        if consultation.estimation_budget:
            tech_parts.append(f"Budget estime: {consultation.estimation_budget}.")
        if consultation.procedure:
            tech_parts.append(f"Procedure: {consultation.procedure}.")
        if consultation.allotissement:
            tech_parts.append(f"Allotissement: {consultation.allotissement}.")
        if consultation.domaines_activite:
            tech_parts.append(f"Domaines d'activite requis: {consultation.domaines_activite}.")
        if consultation.qualifications:
            tech_parts.append(f"Qualifications exigees: {consultation.qualifications}.")

        # Domain-specific technical descriptions
        domain_tech = {
            'dev': (
                "Volet Developpement/Logiciel: Le projet implique des activites de conception, "
                "developpement et/ou maintenance de solutions logicielles. Cela peut inclure "
                "l'analyse des besoins, la redaction de specifications techniques, l'architecture "
                "applicative, le developpement (front-end/back-end), les tests unitaires et "
                "d'integration, le deploiement, et la maintenance corrective/evolutive. "
                "Technologies potentielles: frameworks web, bases de donnees, APIs REST, "
                "methodologie Agile/Scrum."
            ),
            'data': (
                "Volet Data/BI: Le projet comporte une composante de gestion, traitement ou "
                "analyse de donnees. Cela peut inclure la conception de bases de donnees, "
                "le developpement de pipelines ETL, la mise en place de tableaux de bord "
                "(dashboards), le reporting automatise, et l'analyse predictive. "
                "Technologies potentielles: SQL/NoSQL, Power BI/Tableau, Python/R, "
                "Data Warehouse, Data Lake."
            ),
            'ai': (
                "Volet Intelligence Artificielle: Le projet integre des composantes d'IA/ML. "
                "Cela peut inclure la collecte et preparation de donnees d'entrainement, "
                "le developpement de modeles (classification, regression, NLP, vision), "
                "l'optimisation et le fine-tuning, le deploiement de modeles en production "
                "(MLOps), et le monitoring des performances. "
                "Technologies potentielles: Python, TensorFlow/PyTorch, scikit-learn, "
                "Hugging Face, APIs OpenAI."
            ),
            'cloud': (
                "Volet Cloud/Infrastructure: Le projet implique la mise en place ou la gestion "
                "d'infrastructure informatique. Cela peut inclure la virtualisation de serveurs, "
                "la migration vers le cloud, la configuration de reseaux, la mise en place "
                "de solutions de sauvegarde/PRA, le monitoring, et la securisation. "
                "Technologies potentielles: AWS/Azure/GCP, VMware, Docker/Kubernetes, "
                "Terraform, Ansible."
            ),
        }
        for d in domains:
            if d in domain_tech:
                tech_parts.append(domain_tech[d])

        if articles:
            arts = [a.strip().lstrip('#').strip() for a in articles.replace(' | ', '\n').split('\n') if a.strip()]
            if arts:
                tech_parts.append("Articles/Lots identifies: " + "; ".join(arts[:10]) + ".")

        consultation.description_technique = "\n".join(tech_parts)

        # --- Description fonctionnelle ---
        fonc_parts = []
        fonc_parts.append(f"Objectif fonctionnel: {objet}.")

        if consultation.acheteur:
            fonc_parts.append(f"Organisme beneficiaire: {consultation.acheteur}.")
        if consultation.lieu_execution:
            fonc_parts.append(f"Lieu d'execution: {consultation.lieu_execution}.")

        # CPS-enriched functional info
        if consultation.estimation_budget:
            fonc_parts.append(f"Enveloppe budgetaire estimee: {consultation.estimation_budget}.")
        if consultation.reservation_pme:
            fonc_parts.append(f"Note: {consultation.reservation_pme}.")

        # Extract key functional info from CPS content
        if consultation.contenu_cps and len(consultation.contenu_cps) > 50:
            # Pull first meaningful paragraph as context
            cps_lines = [l.strip() for l in consultation.contenu_cps.split('\n')
                         if len(l.strip()) > 30 and not l.strip().startswith(('http', 'www'))]
            if cps_lines:
                fonc_parts.append("Contexte extrait du cahier des charges:")
                for line in cps_lines[:5]:  # Max 5 key lines
                    fonc_parts.append(f"- {line[:200]}")

        domain_fonc = {
            'dev': (
                "Besoins fonctionnels lies au developpement: "
                "- Mise a disposition d'une solution logicielle repondant aux besoins metier de l'acheteur. "
                "- Interface utilisateur ergonomique et accessible. "
                "- Integration avec les systemes existants. "
                "- Formation des utilisateurs finaux. "
                "- Documentation technique et fonctionnelle. "
                "- Garantie et support post-livraison."
            ),
            'data': (
                "Besoins fonctionnels lies aux donnees: "
                "- Centralisation et structuration des donnees de l'organisme. "
                "- Automatisation de la collecte, du nettoyage et de la transformation des donnees. "
                "- Visualisation et suivi d'indicateurs cles (KPI) via des tableaux de bord interactifs. "
                "- Rapports periodiques et exports personnalises. "
                "- Respect des normes de securite et de confidentialite des donnees."
            ),
            'ai': (
                "Besoins fonctionnels lies a l'IA: "
                "- Automatisation de processus manuels repetitifs. "
                "- Aide a la decision basee sur l'analyse predictive. "
                "- Extraction d'informations a partir de documents non structures. "
                "- Amelioration continue des modeles par apprentissage. "
                "- Interface de restitution des resultats comprehensible par les metiers."
            ),
            'cloud': (
                "Besoins fonctionnels lies a l'infrastructure: "
                "- Haute disponibilite et continuite de service. "
                "- Scalabilite pour absorber les pics de charge. "
                "- Sauvegarde automatisee et plan de reprise d'activite (PRA). "
                "- Securisation des acces et des communications. "
                "- Monitoring et alerting proactif."
            ),
        }
        for d in domains:
            if d in domain_fonc:
                fonc_parts.append(domain_fonc[d])

        consultation.description_fonctionnelle = "\n".join(fonc_parts)

        # --- Requirements ---
        reqs = []
        # Common requirements
        reqs.append("Justificatifs d'experience dans des projets similaires (3 references minimum)")
        reqs.append("Equipe projet qualifiee avec CV des intervenants cles")

        if 'dev' in domains:
            reqs.extend([
                "Maitrise des langages/frameworks demandes (web, mobile, ERP/CRM)",
                "Methodologie de gestion de projet (Agile, V-cycle, etc.)",
                "Plan de tests et de recette",
                "Plan de formation utilisateurs",
                "Clause de maintenance et SLA (accord de niveau de service)",
            ])
        if 'data' in domains:
            reqs.extend([
                "Competences en architecture de donnees (SQL, NoSQL, Data Warehouse)",
                "Maitrise des outils BI (Power BI, Tableau, Qlik, etc.)",
                "Experience en ETL et integration de donnees",
                "Conformite aux normes de protection des donnees",
            ])
        if 'ai' in domains:
            reqs.extend([
                "Competences en Machine Learning / Deep Learning",
                "Experience en deploiement de modeles en production (MLOps)",
                "Capacite d'explication des modeles (interpretabilite)",
                "Protocole d'evaluation et de validation des modeles",
            ])
        if 'cloud' in domains:
            reqs.extend([
                "Certifications cloud (AWS, Azure, GCP) ou virtualisation",
                "Experience en migration d'infrastructure",
                "Maitrise des outils de monitoring et securite",
                "Plan PRA/PCA (reprise et continuite d'activite)",
            ])

        # Category-based requirements
        if cat == 'Fournitures':
            reqs.append("Fiches techniques detaillees des equipements/licences proposes")
            reqs.append("Delais de livraison et garantie constructeur")
        elif cat == 'Services':
            reqs.append("Planning previsionnel detaille des prestations")
            reqs.append("Engagement de niveaux de service (SLA)")

        reqs.append("Offre financiere detaillee (bordereau des prix - detail estimatif)")

        consultation.requirements = reqs

    def generate_recommendations(self, consultation: Consultation) -> None:
        """Generate strengths, risks, and recommendations."""
        strengths = []
        risks = []
        recommendations = []

        # Strengths based on domain matches
        domain_labels = {
            'dev': 'Developpement logiciel',
            'data': 'Data et Business Intelligence',
            'ai': 'Intelligence Artificielle',
            'cloud': 'Cloud et Infrastructure'
        }
        for d in consultation.matched_domains:
            strengths.append(f"Correspond au domaine {domain_labels.get(d, d)} d'Alexsys")

        if consultation.score_total >= 6:
            strengths.append("Score tres eleve - alignement fort avec notre expertise")
        if len(consultation.matched_domains) >= 2:
            strengths.append(f"Multi-domaines ({len(consultation.matched_domains)}) - opportunite transversale")

        # Categorie-based strengths
        if consultation.categorie == 'Services':
            strengths.append("Categorie Services - marges potentiellement interessantes")
        elif consultation.categorie == 'Fournitures':
            strengths.append("Fournitures IT - possible vente de licences/equipements")

        # Risk assessment
        if not consultation.date_limite:
            risks.append("Date limite non connue - risque de delai")
        elif consultation.date_limite:
            try:
                dl = datetime.strptime(consultation.date_limite, '%d/%m/%Y')
                days_left = (dl - datetime.now()).days
                if days_left < 5:
                    risks.append(f"URGENT: seulement {days_left} jours restants avant la date limite")
                elif days_left < 14:
                    risks.append(f"Delai court: {days_left} jours avant cloture")
            except ValueError:
                pass

        if not consultation.acheteur:
            risks.append("Acheteur non identifie - difficulte de contact")
        if not consultation.budget_estime and not consultation.estimation_budget:
            risks.append("Budget non communique - estimation difficile")
        if consultation.estimation_budget:
            strengths.append(f"Budget identifie: {consultation.estimation_budget}")
        if consultation.caution_provisoire:
            recommendations.append(f"Preparer la caution provisoire: {consultation.caution_provisoire}")
        if consultation.categorie == 'Travaux':
            risks.append("Categorie Travaux - peut impliquer des composantes non-IT")

        # Recommendations
        if consultation.priority == Priority.HOT:
            recommendations.append("ACTION PRIORITAIRE - Preparer une reponse complete")
            recommendations.append("Telecharger immediatement le cahier des charges (CPS/RC)")
            recommendations.append("Constituer l'equipe projet et estimer les couts")
        elif consultation.priority == Priority.WARM:
            recommendations.append("Analyser le cahier des charges en detail")
            recommendations.append("Evaluer la pertinence par rapport au planning actuel")

        if 'dev' in consultation.matched_domains:
            recommendations.append("Preparer les references projets similaires (dev/integration)")
        if 'data' in consultation.matched_domains:
            recommendations.append("Mettre en avant l'expertise BI/Data d'Alexsys")
        if 'cloud' in consultation.matched_domains:
            recommendations.append("Inclure les certifications cloud et partenariats")
        if 'ai' in consultation.matched_domains:
            recommendations.append("Souligner les competences IA/ML et cas d'usage concrets")

        if consultation.email_contact:
            recommendations.append(f"Contacter l'acheteur: {consultation.email_contact}")
        if consultation.fichiers_joints:
            recommendations.append(f"Telecharger les pieces: {consultation.fichiers_joints}")

        consultation.strengths = strengths
        consultation.risks = risks
        consultation.recommendations = recommendations


# ============================================================
# Document generator (DOCX)
# ============================================================

class DossierGenerator:
    """Generate Word documents for qualified consultations."""

    COMPANY_NAME = "Alexsys Solutions"
    COMPANY_TAGLINE = "Data | BI | IA | Cloud | Developpement"

    def __init__(self, output_dir: str = "dossiers_generes"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)

    def _add_markdown_text(self, doc, text: str):
        """Parse markdown-like text from LLM and add it as formatted DOCX paragraphs."""
        from docx.shared import Pt, RGBColor
        if not text:
            return
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
            # Bold headers (** or ##)
            if line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line.strip('* '))
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0, 51, 102)
            elif line.startswith('## '):
                doc.add_heading(line.lstrip('# '), level=2)
            elif line.startswith('### '):
                doc.add_heading(line.lstrip('# '), level=3)
            # Bullet points
            elif line.startswith(('- ', '• ', '* ')) and not line.startswith('**'):
                bullet_text = line.lstrip('-•* ').strip()
                # Handle inline bold in bullets
                if '**' in bullet_text:
                    p = doc.add_paragraph(style='List Bullet')
                    parts = re.split(r'(\*\*[^*]+\*\*)', bullet_text)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part.strip('* '))
                            run.bold = True
                        else:
                            p.add_run(part)
                else:
                    doc.add_paragraph(bullet_text, style='List Bullet')
            # Numbered items
            elif re.match(r'^\d+\.\s', line):
                text_content = re.sub(r'^\d+\.\s*', '', line)
                if '**' in text_content:
                    p = doc.add_paragraph(style='List Number')
                    parts = re.split(r'(\*\*[^*]+\*\*)', text_content)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part.strip('* '))
                            run.bold = True
                        else:
                            p.add_run(part)
                else:
                    doc.add_paragraph(text_content, style='List Number')
            # Regular paragraph with potential inline bold
            elif '**' in line:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*[^*]+\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part.strip('* '))
                        run.bold = True
                    else:
                        p.add_run(part)
            else:
                doc.add_paragraph(line)

    def generate(self, consultation: Consultation) -> List[str]:
        """Generate both technical and administrative dossiers. Returns list of file paths."""
        ref_clean = re.sub(r'[^\w\-]', '_', consultation.reference or consultation.id)[:30]
        dossier_dir = self.output_dir / ref_clean
        dossier_dir.mkdir(exist_ok=True)

        paths = []
        # 1. Dossier Technique
        try:
            p = self._generate_technique(consultation, dossier_dir, ref_clean)
            if p:
                paths.append(p)
        except Exception as e:
            logger.error(f"Erreur dossier technique {consultation.reference}: {e}")

        # 2. Dossier Administratif
        try:
            p = self._generate_administratif(consultation, dossier_dir, ref_clean)
            if p:
                paths.append(p)
        except Exception as e:
            logger.error(f"Erreur dossier administratif {consultation.reference}: {e}")

        # 3. Save analysis JSON
        self._save_analysis_json(consultation, dossier_dir, ref_clean)

        return paths

    def _make_cover_page(self, doc, consultation: Consultation, doc_type: str):
        """Create a cover page for a dossier."""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        for _ in range(3):
            doc.add_paragraph()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(self.COMPANY_NAME)
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0, 51, 102)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(self.COMPANY_TAGLINE)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(100, 100, 100)

        doc.add_paragraph()
        doc.add_paragraph()

        dt = doc.add_paragraph()
        dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = dt.add_run(doc_type)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph()

        ref_p = doc.add_paragraph()
        ref_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ref_p.add_run(f"Reference: {consultation.reference}").font.size = Pt(14)

        obj_p = doc.add_paragraph()
        obj_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = obj_p.add_run(consultation.objet[:150])
        run.font.size = Pt(12)
        run.italic = True

        doc.add_paragraph()

        info_p = doc.add_paragraph()
        info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_p.add_run(f"Acheteur: {consultation.acheteur or 'Non specifie'}").font.size = Pt(11)

        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_p.add_run(f"Date: {datetime.now().strftime('%d/%m/%Y')}")

        doc.add_page_break()

    def _add_footer(self, doc, doc_type: str):
        """Add footer to document."""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc.add_paragraph()
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run(
            f"{doc_type} - Genere par Alexsys AI Platform - {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(150, 150, 150)

    # ================================================================
    # DOSSIER TECHNIQUE
    # ================================================================
    def _generate_technique(self, c: Consultation, dossier_dir: Path, ref_clean: str) -> str:
        """Generate the technical dossier DOCX."""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Cover
        self._make_cover_page(doc, c, "DOSSIER TECHNIQUE")

        # Sommaire
        doc.add_heading("SOMMAIRE", level=1)
        sections = [
            "1. Fiche de synthese",
            "2. Comprehension du besoin",
            "3. Description technique",
            "4. Description fonctionnelle",
            "5. Exigences et prerequis",
            "6. Analyse strategique",
            "7. Methodologie et approche",
            "8. Architecture proposee",
            "9. Equipe projet",
            "10. Planning previsionnel",
            "11. Livrables",
            "12. Garantie et maintenance",
        ]
        for s in sections:
            p = doc.add_paragraph(s)
            p.paragraph_format.space_after = Pt(4)
        doc.add_page_break()

        # --- 1. Fiche de synthese ---
        doc.add_heading("1. Fiche de synthese", level=1)
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        synth = [
            ("Reference", c.reference),
            ("Objet", c.objet[:200]),
            ("Acheteur", c.acheteur or "Non specifie"),
            ("Categorie", c.categorie),
            ("Nature prestation", c.nature_prestation or "-"),
            ("Lieu d'execution", c.lieu_execution or "-"),
            ("Date limite", c.date_limite),
            ("Budget estime", c.estimation_budget or c.budget_estime or "Non communique"),
            ("Caution provisoire", c.caution_provisoire or "-"),
            ("Procedure", c.procedure or "-"),
            ("Priorite Alexsys", c.priority.value),
            ("Score IT", f"{c.score_total}"),
            ("Domaines", ", ".join(c.matched_domains) if c.matched_domains else "-"),
            ("Confiance NLP", f"{c.nlp_confidence:.0%}" if c.nlp_confidence else "-"),
            ("Source CPS", c.cps_source or "Non extraite"),
        ]
        for label, value in synth:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value)
            for p in row.cells[0].paragraphs:
                for r in p.runs:
                    r.bold = True
        doc.add_paragraph()

        # --- 2. Comprehension du besoin ---
        doc.add_heading("2. Comprehension du besoin", level=1)
        doc.add_paragraph(f"L'objet de la consultation est: {c.objet}")
        if c.nature_prestation:
            doc.add_paragraph(f"Nature de la prestation attendue: {c.nature_prestation}")
        if c.articles:
            doc.add_heading("Articles / Lots", level=2)
            for art in c.articles.replace(' | ', '\n').replace('#', '').split('\n'):
                art = art.strip()
                if art:
                    doc.add_paragraph(art, style='List Bullet')
        if c.nlp_key_entities:
            doc.add_heading("Entites cles identifiees (analyse NLP)", level=2)
            for ent in c.nlp_key_entities:
                doc.add_paragraph(ent, style='List Bullet')

        # CPS-enriched information
        if c.domaines_activite or c.qualifications:
            doc.add_heading("Informations extraites du portail", level=2)
            if c.domaines_activite:
                doc.add_paragraph(f"Domaines d'activite requis: {c.domaines_activite}")
            if c.qualifications:
                doc.add_paragraph(f"Qualifications exigees: {c.qualifications}")
            if c.allotissement:
                doc.add_paragraph(f"Allotissement: {c.allotissement}")
            if c.reservation_pme:
                doc.add_paragraph(f"Reservation: {c.reservation_pme}")

        # CPS content extract
        if c.contenu_cps and len(c.contenu_cps) > 50:
            doc.add_heading("Extrait du cahier des charges", level=2)
            doc.add_paragraph(f"(Source: {c.cps_source or 'portail'})")
            cps_lines = [l.strip() for l in c.contenu_cps.split('\n')
                         if len(l.strip()) > 15]
            for line in cps_lines[:30]:  # Max 30 lines in dossier
                doc.add_paragraph(line[:300])

        # Score details
        doc.add_heading("Pertinence par domaine", level=2)
        domain_labels = {'dev': 'Dev/Logiciel', 'data': 'Data/BI', 'ai': 'IA/ML', 'cloud': 'Cloud/Infra'}
        st = doc.add_table(rows=0, cols=2)
        st.style = 'Light Grid Accent 1'
        for domain, score in c.domain_scores.items():
            if score > 0:
                row = st.add_row()
                row.cells[0].text = domain_labels.get(domain, domain)
                row.cells[1].text = "=" * score + f" ({score})"
        doc.add_paragraph()

        # --- 3. Description technique ---
        doc.add_heading("3. Description technique", level=1)
        if c.description_technique:
            self._add_markdown_text(doc, c.description_technique)
        else:
            doc.add_paragraph("Description technique a completer apres analyse du CPS.")

        # --- 4. Description fonctionnelle ---
        doc.add_heading("4. Description fonctionnelle", level=1)
        if c.description_fonctionnelle:
            self._add_markdown_text(doc, c.description_fonctionnelle)
        else:
            doc.add_paragraph("Description fonctionnelle a completer apres analyse du CPS.")

        # --- 5. Exigences et prerequis ---
        doc.add_heading("5. Exigences et prerequis", level=1)
        if c.requirements:
            for req in c.requirements:
                doc.add_paragraph(req, style='List Bullet')
        else:
            doc.add_paragraph("Les exigences seront detaillees apres analyse du cahier des charges.")

        # --- 6. Analyse strategique ---
        doc.add_heading("6. Analyse strategique", level=1)
        has_strategic = c.strengths or c.risks or c.recommendations
        if has_strategic:
            if c.strengths:
                doc.add_heading("6.1 Points forts de notre candidature", level=2)
                for s in c.strengths:
                    doc.add_paragraph(s, style='List Bullet')
            if c.risks:
                doc.add_heading("6.2 Risques identifies et mesures d'attenuation", level=2)
                for r in c.risks:
                    doc.add_paragraph(r, style='List Bullet')
            if c.recommendations:
                doc.add_heading("6.3 Recommandations et actions prioritaires", level=2)
                for rec in c.recommendations:
                    doc.add_paragraph(rec, style='List Bullet')
        else:
            doc.add_paragraph(
                "L'analyse strategique detaillee sera elaboree apres examen approfondi "
                "du cahier des charges et du reglement de consultation."
            )

        # --- 7. Methodologie et approche ---
        doc.add_heading("7. Methodologie et approche", level=1)

        if 'dev' in c.matched_domains:
            doc.add_heading("6.1 Approche Developpement", level=2)
            doc.add_paragraph(
                "Notre methodologie de developpement repose sur les pratiques Agile/Scrum:\n"
                "- Sprint planning: Definition des objectifs et du backlog par iteration (2-3 semaines)\n"
                "- Developpement iteratif: Conception, codage, tests unitaires, revue de code\n"
                "- Integration continue (CI/CD): Automatisation des builds, tests et deploiements\n"
                "- Revues de sprint: Demonstration des fonctionnalites, feedback client\n"
                "- Retrospectives: Amelioration continue du processus\n"
                "- Recette: Tests d'acceptation utilisateur (UAT) et validation fonctionnelle"
            )
        if 'data' in c.matched_domains:
            doc.add_heading("6.2 Approche Data/BI", level=2)
            doc.add_paragraph(
                "Notre demarche Data suit un processus structure:\n"
                "- Audit de l'existant: Inventaire des sources de donnees, qualite, volumetrie\n"
                "- Conception: Modelisation du Data Warehouse/Data Lake, schema en etoile/flocon\n"
                "- ETL: Developpement des flux d'extraction, transformation, chargement\n"
                "- Visualisation: Conception des tableaux de bord et rapports interactifs\n"
                "- Validation: Tests de coherence, performance, et acceptation utilisateur\n"
                "- Formation: Accompagnement des utilisateurs et transfert de competences"
            )
        if 'ai' in c.matched_domains:
            doc.add_heading("6.3 Approche IA/ML", level=2)
            doc.add_paragraph(
                "Notre methodologie IA suit le cycle CRISP-DM:\n"
                "- Comprehension metier: Definition des objectifs et KPI de succes\n"
                "- Comprehension des donnees: Exploration, analyse statistique, qualite\n"
                "- Preparation des donnees: Nettoyage, feature engineering, augmentation\n"
                "- Modelisation: Selection d'algorithmes, entrainement, cross-validation\n"
                "- Evaluation: Metriques de performance, tests A/B, interpretabilite\n"
                "- Deploiement: Mise en production (MLOps), monitoring, retraining automatique"
            )
        if 'cloud' in c.matched_domains:
            doc.add_heading("6.4 Approche Cloud/Infrastructure", level=2)
            doc.add_paragraph(
                "Notre demarche Infrastructure suit les bonnes pratiques:\n"
                "- Audit: Cartographie de l'infrastructure existante, identification des gaps\n"
                "- Conception: Architecture cible, dimensionnement, choix technologiques\n"
                "- Migration: Plan de migration par phases, tests de non-regression\n"
                "- Securisation: Hardening, firewall, chiffrement, gestion des acces (IAM)\n"
                "- Monitoring: Mise en place d'outils de supervision et d'alerting\n"
                "- Documentation: Procedures d'exploitation, PRA/PCA"
            )

        if not c.matched_domains:
            doc.add_paragraph(
                "La methodologie detaillee sera proposee apres analyse approfondie du cahier des charges."
            )

        # --- 8. Architecture proposee ---
        doc.add_heading("8. Architecture proposee", level=1)
        doc.add_paragraph(
            "L'architecture technique sera detaillee dans la reponse complete, "
            "en fonction des exigences specifiques du cahier des charges. "
            "Elle respectera les principes suivants:"
        )
        archi_principles = [
            "Separation des couches (presentation, logique metier, donnees)",
            "Haute disponibilite et tolerance aux pannes",
            "Securite by design (authentification, autorisation, chiffrement)",
            "Scalabilite horizontale et verticale",
            "Interoperabilite et standards ouverts",
            "Conformite aux normes et reglementations en vigueur",
        ]
        for ap in archi_principles:
            doc.add_paragraph(ap, style='List Bullet')

        # --- 9. Equipe projet ---
        doc.add_heading("9. Equipe projet", level=1)
        doc.add_paragraph(
            "L'equipe projet proposee sera composee de profils qualifies et experimentes, "
            "adaptes aux besoins specifiques de la consultation:"
        )
        team_table = doc.add_table(rows=1, cols=3)
        team_table.style = 'Light Grid Accent 1'
        hdr = team_table.rows[0].cells
        hdr[0].text = "Role"
        hdr[1].text = "Profil"
        hdr[2].text = "Experience"
        for p in hdr[0].paragraphs:
            for r in p.runs:
                r.bold = True
        for p in hdr[1].paragraphs:
            for r in p.runs:
                r.bold = True
        for p in hdr[2].paragraphs:
            for r in p.runs:
                r.bold = True

        team_roles = [("Chef de projet", "Senior, certifie PMP/PRINCE2", "10+ ans")]
        if 'dev' in c.matched_domains:
            team_roles.extend([
                ("Architecte logiciel", "Expert architecture applicative", "8+ ans"),
                ("Developpeur senior", "Full-stack (front-end + back-end)", "5+ ans"),
                ("Testeur/QA", "Specialiste tests automatises", "3+ ans"),
            ])
        if 'data' in c.matched_domains:
            team_roles.extend([
                ("Data Engineer", "Expert ETL et architecture donnees", "5+ ans"),
                ("Analyste BI", "Specialiste reporting et dashboards", "4+ ans"),
            ])
        if 'ai' in c.matched_domains:
            team_roles.extend([
                ("Data Scientist", "Expert ML/DL, Python, R", "5+ ans"),
                ("ML Engineer", "Specialiste MLOps et deploiement", "4+ ans"),
            ])
        if 'cloud' in c.matched_domains:
            team_roles.extend([
                ("Architecte Cloud", "Certifie AWS/Azure/GCP", "6+ ans"),
                ("Ingenieur systeme", "Expert virtualisation, securite", "5+ ans"),
            ])

        for role, profil, exp in team_roles:
            row = team_table.add_row()
            row.cells[0].text = role
            row.cells[1].text = profil
            row.cells[2].text = exp
        doc.add_paragraph()

        # --- 10. Planning previsionnel ---
        doc.add_heading("10. Planning previsionnel", level=1)
        doc.add_paragraph(
            "Le planning detaille sera ajuste en fonction du delai contractuel "
            "et de la complexite reelle du projet. Phases indicatives:"
        )
        phases = [
            ("Phase 1 - Cadrage", "Reunion de lancement, analyse detaillee du besoin, validation du perimetre"),
            ("Phase 2 - Conception", "Architecture technique et fonctionnelle, specifications detaillees"),
            ("Phase 3 - Realisation", "Developpement/integration, tests unitaires, integration continue"),
            ("Phase 4 - Recette", "Tests d'acceptation, corrections, validation utilisateur"),
            ("Phase 5 - Deploiement", "Mise en production, formation, documentation"),
            ("Phase 6 - Garantie", "Support post-livraison, maintenance corrective, suivi"),
        ]
        plan_table = doc.add_table(rows=1, cols=2)
        plan_table.style = 'Light Grid Accent 1'
        ph = plan_table.rows[0].cells
        ph[0].text = "Phase"
        ph[1].text = "Activites"
        for p in ph[0].paragraphs:
            for r in p.runs:
                r.bold = True
        for p in ph[1].paragraphs:
            for r in p.runs:
                r.bold = True
        for phase, desc in phases:
            row = plan_table.add_row()
            row.cells[0].text = phase
            row.cells[1].text = desc
        doc.add_paragraph()

        # --- 11. Livrables ---
        doc.add_heading("11. Livrables", level=1)
        livrables = [
            "Document d'analyse et de specification detaillee",
            "Document d'architecture technique",
            "Code source commente et versionne (Git)",
            "Jeux de tests et rapports de tests",
            "Documentation technique (installation, exploitation, API)",
            "Documentation utilisateur (guide, FAQ)",
            "Support de formation",
            "Proces-verbal de recette",
        ]
        if 'data' in c.matched_domains:
            livrables.extend([
                "Modele de donnees (MCD/MLD)",
                "Specifications des flux ETL",
                "Tableaux de bord et rapports BI",
            ])
        if 'ai' in c.matched_domains:
            livrables.extend([
                "Rapport d'analyse exploratoire des donnees",
                "Modeles entraines et metriques de performance",
                "Pipeline MLOps documente",
            ])
        if 'cloud' in c.matched_domains:
            livrables.extend([
                "Dossier d'architecture infrastructure",
                "Plan de migration",
                "Procedures d'exploitation et PRA/PCA",
            ])
        for liv in livrables:
            doc.add_paragraph(liv, style='List Bullet')

        # --- 12. Garantie et maintenance ---
        doc.add_heading("12. Garantie et maintenance", level=1)
        doc.add_paragraph(
            "Alexsys Solutions propose une periode de garantie apres la recette definitive, "
            "couvrant:"
        )
        garantie = [
            "Correction des anomalies et bugs identifies",
            "Support technique par email, telephone et ticketing",
            "Mises a jour de securite et patches critiques",
            "Assistance aux utilisateurs pendant la phase de demarrage",
        ]
        for g in garantie:
            doc.add_paragraph(g, style='List Bullet')
        doc.add_paragraph(
            "Au-dela de la garantie, un contrat de maintenance evolutive et corrective "
            "pourra etre propose avec des niveaux de service (SLA) adaptes."
        )

        self._add_footer(doc, "Dossier Technique")

        filename = f"dossier_technique_{ref_clean}_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = dossier_dir / filename
        doc.save(str(filepath))
        logger.info(f"Dossier technique: {filepath}")
        return str(filepath)

    # ================================================================
    # DOSSIER ADMINISTRATIF
    # ================================================================
    def _generate_administratif(self, c: Consultation, dossier_dir: Path, ref_clean: str) -> str:
        """Generate the administrative dossier DOCX."""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Cover
        self._make_cover_page(doc, c, "DOSSIER ADMINISTRATIF")

        # Sommaire
        doc.add_heading("SOMMAIRE", level=1)
        sections = [
            "1. Identification de la consultation",
            "2. Presentation de la societe",
            "3. Moyens humains et materiels",
            "4. References et experiences",
            "5. Certifications et agrements",
            "6. Pieces administratives",
            "7. Engagement et declarations",
            "8. Recommandations et actions",
        ]
        for s in sections:
            p = doc.add_paragraph(s)
            p.paragraph_format.space_after = Pt(4)
        doc.add_page_break()

        # --- 1. Identification consultation ---
        doc.add_heading("1. Identification de la consultation", level=1)
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        info = [
            ("Reference", c.reference),
            ("Objet", c.objet[:200]),
            ("Acheteur", c.acheteur or "Non specifie"),
            ("Categorie", c.categorie),
            ("Nature prestation", c.nature_prestation or "-"),
            ("Lieu d'execution", c.lieu_execution or "-"),
            ("Date publication", c.date_publication),
            ("Date limite de depot", c.date_limite),
            ("Budget estime", c.estimation_budget or c.budget_estime or "Non communique"),
            ("Caution provisoire", c.caution_provisoire or "-"),
            ("Procedure", c.procedure or "-"),
            ("Source", c.source),
            ("URL", c.url),
        ]
        for label, value in info:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value)
            for p in row.cells[0].paragraphs:
                for r in p.runs:
                    r.bold = True

        if c.fichiers_joints:
            doc.add_paragraph()
            doc.add_paragraph(f"Fichiers joints: {c.fichiers_joints}")
        if c.email_contact:
            doc.add_paragraph(f"Contact: {c.email_contact}")
        doc.add_paragraph()

        # --- 2. Presentation societe ---
        doc.add_heading("2. Presentation de la societe", level=1)

        doc.add_heading("2.1 Identification", level=2)
        soc_table = doc.add_table(rows=0, cols=2)
        soc_table.style = 'Light Grid Accent 1'
        soc_info = [
            ("Raison sociale", "Alexsys Solutions"),
            ("Siege social", "37 Allee des Eucalyptus, Ain-Sebaa, 20590, Casablanca"),
        ]
        for label, value in soc_info:
            row = soc_table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value
            for p in row.cells[0].paragraphs:
                for r in p.runs:
                    r.bold = True
        doc.add_paragraph()

        doc.add_heading("2.2 Domaines d'activite", level=2)
        doc.add_paragraph(
            "Alexsys Solutions est une societe marocaine specialisee dans les technologies "
            "de l'information et le conseil numerique. Nos domaines d'expertise couvrent:"
        )
        expertise = [
            "Developpement logiciel sur mesure (Web, Mobile, ERP, CRM)",
            "Business Intelligence, Data Analytics et Big Data",
            "Intelligence Artificielle et Machine Learning",
            "Infrastructure Cloud, Virtualisation et DevOps",
            "Cybersecurite, Audit et Conformite",
            "Transformation numerique et Infogerance",
            "Formation et accompagnement au changement",
        ]
        for item in expertise:
            doc.add_paragraph(item, style='List Bullet')

        doc.add_paragraph()
        doc.add_paragraph(
            "Notre positionnement multi-domaines nous permet d'offrir des solutions "
            "integrees et coherentes, adaptees aux besoins specifiques de chaque client, "
            "qu'il soit public ou prive."
        )

        # --- 3. Moyens humains et materiels ---
        doc.add_heading("3. Moyens humains et materiels", level=1)

        doc.add_heading("3.1 Moyens humains", level=2)
        doc.add_paragraph(
            "Alexsys Solutions dispose d'une equipe multidisciplinaire composee de:"
        )
        rh = [
            "Chefs de projet seniors certifies (PMP, PRINCE2, Scrum Master)",
            "Architectes logiciels et systeme",
            "Developpeurs full-stack (Java, Python, .NET, JavaScript/TypeScript)",
            "Data Engineers et Data Scientists",
            "Experts BI (Power BI, Tableau, Qlik)",
            "Ingenieurs Cloud et systeme (AWS, Azure, VMware)",
            "Specialistes cybersecurite (CEH, CISSP)",
            "Consultants fonctionnels et formateurs",
        ]
        for item in rh:
            doc.add_paragraph(item, style='List Bullet')

        doc.add_heading("3.2 Moyens materiels et techniques", level=2)
        materiels = [
            "Environnements de developpement et de test dedies",
            "Plateforme d'integration continue (CI/CD) - GitLab/Jenkins",
            "Infrastructure de demonstration et de recette",
            "Outils de gestion de projet (Jira, Confluence, MS Project)",
            "Outils de tests automatises (Selenium, JMeter, SonarQube)",
            "Licences logicielles professionnelles a jour",
        ]
        for item in materiels:
            doc.add_paragraph(item, style='List Bullet')

        # --- 4. References ---
        doc.add_heading("4. References et experiences", level=1)
        doc.add_paragraph(
            "Alexsys Solutions a realise de nombreux projets pour des clients de reference "
            "dans les secteurs public et prive."
        )

        # --- 5. Certifications ---
        doc.add_heading("5. Certifications et agrements", level=1)
        doc.add_paragraph("Certifications et agrements detenus ou en cours:")
        certs = [
            "ISO 9001 - Systeme de management de la qualite",
            "ISO 27001 - Securite de l'information",
            "Partenaire Microsoft (Silver/Gold)",
            "Partenaire Oracle / SAP (si applicable)",
            "Certification PMI/PMP pour les chefs de projet",
            "Certifications cloud (AWS Solutions Architect, Azure Administrator, etc.)",
            "CMMI niveau X (si applicable)",
        ]
        for cert in certs:
            doc.add_paragraph(cert, style='List Bullet')

        # --- 6. Pieces administratives ---
        doc.add_heading("6. Pieces administratives", level=1)
        doc.add_paragraph(
            "Conformement au reglement de consultation et a la reglementation en vigueur "
            "(Decret n 2-12-349 relatif aux marches publics), les pieces suivantes "
            "sont requises pour la constitution du dossier administratif:"
        )

        doc.add_heading("6.1 Dossier administratif", level=2)
        admin_pieces = [
            ("Lettre d'accompagnement", "A preparer - signer et cacheter"),
            ("Declaration sur l'honneur", "Modele RC - signer et cacheter"),
            ("Certificat d'immatriculation au registre de commerce", "Original ou copie certifiee conforme"),
            ("Attestation de regularite fiscale", "En cours de validite (< 6 mois)"),
            ("Attestation CNSS", "En cours de validite (< 6 mois)"),
            ("Caution provisoire", "Selon montant indique dans le RC (si applicable)"),
            ("Pouvoir / Habilitation du signataire", "Si le signataire n'est pas le representant legal"),
            ("Statuts de la societe", "Copie certifiee conforme"),
            ("Certificat de nationalite", "Si exige par le RC"),
        ]
        adm_table = doc.add_table(rows=1, cols=2)
        adm_table.style = 'Light Grid Accent 1'
        adm_table.rows[0].cells[0].text = "Piece"
        adm_table.rows[0].cells[1].text = "Statut / Remarque"
        for p in adm_table.rows[0].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
        for p in adm_table.rows[0].cells[1].paragraphs:
            for r in p.runs:
                r.bold = True
        for piece, statut in admin_pieces:
            row = adm_table.add_row()
            row.cells[0].text = piece
            row.cells[1].text = statut
        doc.add_paragraph()

        doc.add_heading("6.2 Dossier technique (a joindre)", level=2)
        tech_pieces = [
            ("Note de presentation de la societe", "Historique, domaines, effectifs"),
            ("References et attestations de bonne execution", "3 references minimum"),
            ("Moyens humains (CV)", "CV des profils cles affectes au projet"),
            ("Offre technique detaillee", "Voir Dossier Technique separe"),
        ]
        tech_table = doc.add_table(rows=1, cols=2)
        tech_table.style = 'Light Grid Accent 1'
        tech_table.rows[0].cells[0].text = "Piece"
        tech_table.rows[0].cells[1].text = "Remarque"
        for p in tech_table.rows[0].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
        for p in tech_table.rows[0].cells[1].paragraphs:
            for r in p.runs:
                r.bold = True
        for piece, note in tech_pieces:
            row = tech_table.add_row()
            row.cells[0].text = piece
            row.cells[1].text = note
        doc.add_paragraph()

        doc.add_heading("6.3 Offre financiere", level=2)
        fin_pieces = [
            "Acte d'engagement (modele fourni dans le RC)",
            "Bordereau des prix - Detail estimatif (BDP-DE)",
            "Decomposition du montant global (si forfait)",
            "Sous-detail des prix (si demande)",
        ]
        for fp in fin_pieces:
            doc.add_paragraph(fp, style='List Bullet')

        # --- 7. Engagement ---
        doc.add_heading("7. Engagement et declarations", level=1)
        doc.add_paragraph(
            "Par la presente, Alexsys Solutions s'engage a:"
        )
        engagements = [
            "Respecter les termes et conditions du reglement de consultation",
            "Fournir les prestations conformement au cahier des prescriptions speciales (CPS)",
            "Respecter les delais contractuels d'execution",
            "Garantir la confidentialite des informations echangees",
            "Mettre a disposition les moyens humains et materiels necessaires",
            "Se conformer a la legislation marocaine en matiere de marches publics",
        ]
        for eng in engagements:
            doc.add_paragraph(eng, style='List Bullet')

        doc.add_paragraph()
        doc.add_paragraph(
            "Fait a ______________, le ___/___/______"
        )
        doc.add_paragraph()
        doc.add_paragraph("Signature et cachet de la societe:")
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph("_" * 40)

        # --- 8. Recommandations ---
        doc.add_heading("8. Recommandations et actions", level=1)
        if c.recommendations:
            doc.add_heading("Actions recommandees", level=2)
            for rec in c.recommendations:
                doc.add_paragraph(rec, style='List Bullet')

        if c.strengths:
            doc.add_heading("Points forts de notre candidature", level=2)
            for s in c.strengths:
                doc.add_paragraph(s, style='List Bullet')

        if c.risks:
            doc.add_heading("Points de vigilance", level=2)
            for r in c.risks:
                doc.add_paragraph(r, style='List Bullet')

        self._add_footer(doc, "Dossier Administratif")

        filename = f"dossier_administratif_{ref_clean}_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = dossier_dir / filename
        doc.save(str(filepath))
        logger.info(f"Dossier administratif: {filepath}")
        return str(filepath)

    def _save_analysis_json(self, c: Consultation, dossier_dir: Path, ref_clean: str):
        """Save the analysis JSON file."""
        analysis = {
            "reference": c.reference,
            "objet": c.objet,
            "acheteur": c.acheteur,
            "priority": c.priority.value,
            "score_total": c.score_total,
            "domain_scores": c.domain_scores,
            "matched_domains": c.matched_domains,
            "matched_keywords": c.matched_keywords,
            "nlp_categories": c.nlp_categories,
            "nlp_confidence": c.nlp_confidence,
            "nlp_key_entities": c.nlp_key_entities,
            "description_technique": c.description_technique,
            "description_fonctionnelle": c.description_fonctionnelle,
            "requirements": c.requirements,
            "strengths": c.strengths,
            "risks": c.risks,
            "recommendations": c.recommendations,
            "estimation_budget": c.estimation_budget,
            "caution_provisoire": c.caution_provisoire,
            "procedure": c.procedure,
            "allotissement": c.allotissement,
            "domaines_activite": c.domaines_activite,
            "qualifications": c.qualifications,
            "reservation_pme": c.reservation_pme,
            "cps_source": c.cps_source,
            "contenu_cps_extrait": c.contenu_cps[:500] if c.contenu_cps else "",
            "url": c.url,
            "date_limite": c.date_limite,
            "generated_at": datetime.now().isoformat(),
        }
        json_path = dossier_dir / f"analyse_{ref_clean}.json"
        with open(json_path, 'w', encoding='utf-8') as jf:
            json.dump(analysis, jf, ensure_ascii=False, indent=2)


# ============================================================
# Pipeline orchestrator
# ============================================================

class Pipeline:
    """
    Orchestrates: Load CSV -> Score -> Enrich CPS -> RAG/NLP Analyse -> Generate Dossiers -> Export

    Usage:
        pipeline = Pipeline()
        results = pipeline.run("data/appels_offres_profond_20260310_100313.csv")
    """

    def __init__(self, output_dir: str = "dossiers_generes", use_nlp: bool = True,
                 enrich_cps: bool = True, use_rag: bool = False):
        self.output_dir = output_dir
        self.use_nlp = use_nlp
        self.use_rag = use_rag
        self.enrich_cps_flag = enrich_cps
        self.scorer = ScoringEngine()
        self.cps_extractor = CPSExtractor() if enrich_cps else None
        self.nlp_analyzer = NLPAnalyzer() if use_nlp else None
        self.rag_engine = None
        if use_rag:
            try:
                from core.rag_engine import RAGEngine
                self.rag_engine = RAGEngine()
                logger.info("RAG engine initialized")
            except Exception as e:
                logger.warning(f"RAG engine init failed: {e}. Falling back to NLP.")
                self.use_rag = False
        self.generator = DossierGenerator(output_dir)
        self.consultations: List[Consultation] = []

    # ----- Stage 1: Load -----
    def load_csv(self, csv_path: str) -> int:
        """Load consultations from a semicolon-delimited CSV."""
        path = Path(csv_path)
        if not path.exists():
            raise FileNotFoundError(f"CSV introuvable: {csv_path}")

        with open(path, 'r', encoding='utf-8-sig') as f:
            reader = csv.DictReader(f, delimiter=';')
            self.consultations = [Consultation.from_csv_row(row) for row in reader]

        logger.info(f"Stage 1 - Charge: {len(self.consultations)} consultations depuis {csv_path}")
        return len(self.consultations)

    # ----- Stage 2: Score -----
    def score_all(self) -> Dict[str, int]:
        """Score all consultations with weighted keyword matching."""
        stats = {"total": 0, "hot": 0, "warm": 0, "cold": 0, "excluded": 0}

        for c in self.consultations:
            self.scorer.score(c)
            stats["total"] += 1
            if c.priority == Priority.HOT:
                stats["hot"] += 1
            elif c.priority == Priority.WARM:
                stats["warm"] += 1
            elif c.priority == Priority.COLD:
                stats["cold"] += 1
            else:
                stats["excluded"] += 1

        logger.info(
            f"Stage 2 - Score: {stats['hot']} HOT, {stats['warm']} WARM, "
            f"{stats['cold']} COLD, {stats['excluded']} exclus"
        )
        return stats

    # ----- Stage 2.5: Enrich from CPS / detail pages -----
    def enrich_from_cps(self, min_priority: Priority = Priority.COLD) -> Dict[str, int]:
        """Enrich IT-relevant consultations with data from detail pages."""
        if not self.cps_extractor:
            return {"enriched": 0, "prado": 0, "bdc": 0, "failed": 0}

        priorities = {Priority.HOT, Priority.WARM}
        if min_priority == Priority.COLD:
            priorities.add(Priority.COLD)

        stats = {"enriched": 0, "prado": 0, "bdc": 0, "failed": 0}
        targets = [c for c in self.consultations if c.priority in priorities]

        for i, c in enumerate(targets, 1):
            portal = "PRADO" if '/bdc/' not in (c.url or '') else "BDC"
            logger.info(f"  Enrichissement {i}/{len(targets)}: {c.reference} ({portal})")
            try:
                ok = self.cps_extractor.enrich(c)
                if ok:
                    stats["enriched"] += 1
                    if c.cps_source == 'prado_html':
                        stats["prado"] += 1
                    elif c.cps_source in ('bdc_pdf', 'bdc_ocr'):
                        stats["bdc"] += 1
                else:
                    stats["failed"] += 1
            except Exception as e:
                logger.warning(f"Enrichissement echoue pour {c.reference}: {e}")
                stats["failed"] += 1

        logger.info(
            f"Stage 2.5 - Enrichissement: {stats['enriched']} enrichis "
            f"(PRADO: {stats['prado']}, BDC: {stats['bdc']}), {stats['failed']} echoues"
        )
        return stats

    # ----- Stage 3: NLP Analysis -----
    def analyze_nlp(self, min_priority: Priority = Priority.COLD) -> int:
        """Run NLP analysis on qualified consultations."""
        priorities_to_analyze = {Priority.HOT, Priority.WARM}
        if min_priority == Priority.COLD:
            priorities_to_analyze.add(Priority.COLD)

        # Create a temporary NLPAnalyzer for descriptions/recommendations
        # (these don't need spaCy)
        desc_analyzer = NLPAnalyzer.__new__(NLPAnalyzer)
        desc_analyzer._nlp = None

        analyzed = 0
        for c in self.consultations:
            if c.priority in priorities_to_analyze:
                # NLP analysis (requires spaCy) - optional
                if self.nlp_analyzer:
                    try:
                        self.nlp_analyzer.analyze(c)
                    except Exception as e:
                        logger.debug(f"NLP spaCy skipped for {c.reference}: {e}")

                # Descriptions & recommendations (no spaCy needed) - always run
                try:
                    desc_analyzer.generate_descriptions(c)
                    desc_analyzer.generate_recommendations(c)
                    analyzed += 1
                except Exception as e:
                    logger.warning(f"Description error for {c.reference}: {e}")

        if not self.nlp_analyzer:
            logger.warning("spaCy non disponible - descriptions generees sans NLP")
        logger.info(f"Stage 3 - NLP: {analyzed} consultations analysees")
        return analyzed

    # ----- Stage 3.5: RAG Enrichment (replaces NLP descriptions with LLM-generated) -----
    def analyze_rag(self, min_priority: Priority = Priority.COLD) -> Dict[str, int]:
        """Run RAG analysis on qualified consultations. Overwrites NLP descriptions."""
        if not self.rag_engine:
            return {"rag_enriched": 0, "rag_failed": 0, "rag_indexed": 0}

        priorities_to_analyze = {Priority.HOT, Priority.WARM}
        if min_priority == Priority.COLD:
            priorities_to_analyze.add(Priority.COLD)

        relevant = [c for c in self.consultations if c.priority in priorities_to_analyze]

        # Step 1: Index all relevant consultations
        logger.info(f"RAG: indexing {len(relevant)} consultations...")
        idx_stats = self.rag_engine.index_all(relevant)

        # Step 2: Generate descriptions for each
        enriched = 0
        failed = 0
        for i, c in enumerate(relevant, 1):
            ref = c.reference or c.id
            try:
                logger.info(f"  RAG {i}/{len(relevant)}: {ref}")
                success = self.rag_engine.enrich_consultation(c)
                if success:
                    enriched += 1
                else:
                    failed += 1
            except Exception as e:
                logger.warning(f"  RAG failed for {ref}: {e}")
                failed += 1

        stats = {
            "rag_enriched": enriched,
            "rag_failed": failed,
            "rag_indexed": idx_stats.get("total_chunks", 0),
        }
        logger.info(f"Stage 3.5 - RAG: {enriched} enrichis, {failed} echoues, {idx_stats.get('total_chunks', 0)} chunks indexes")
        return stats

    # ----- Stage 4: Generate Dossiers -----
    def generate_dossiers(self, min_priority: Priority = Priority.WARM) -> List[str]:
        """Generate DOCX dossiers for qualified consultations."""
        priorities = {Priority.HOT}
        if min_priority == Priority.WARM:
            priorities.add(Priority.WARM)
        elif min_priority == Priority.COLD:
            priorities.update({Priority.WARM, Priority.COLD})

        generated = []
        for c in self.consultations:
            if c.priority in priorities:
                try:
                    paths = self.generator.generate(c)
                    if paths:
                        generated.extend(paths)
                except Exception as e:
                    logger.error(f"Erreur generation dossier {c.reference}: {e}")

        logger.info(f"Stage 4 - Generation: {len(generated)} dossiers crees dans {self.output_dir}/")
        return generated

    # ----- Stage 5: Export results -----
    def export_results(self, output_path: Optional[str] = None) -> str:
        """Export scored consultations to CSV + Excel with user-requested columns."""
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        # By default, overwrite stable "latest" exports to avoid accumulating timestamped files.
        # Set PIPELINE_TIMESTAMP_EXPORTS=1 to keep timestamped filenames.
        keep_ts = os.environ.get("PIPELINE_TIMESTAMP_EXPORTS", "0").strip() == "1"
        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        if not output_path:
            output_path = f"data/pipeline_results_{ts}" if keep_ts else "data/pipeline_results_latest"
        base = output_path.replace('.xlsx', '').replace('.csv', '')
        csv_path = f"{base}.csv"
        xlsx_path = f"{base}.xlsx"

        # Export all IT consultations (including EXCLUDED) so downstream dashboards can decide what to show.
        priority_order = {
            Priority.HOT: 0,
            Priority.WARM: 1,
            Priority.COLD: 2,
            Priority.EXCLUDED: 3,
        }
        relevant = list(self.consultations)
        relevant.sort(key=lambda c: (priority_order.get(c.priority, 9), -c.score_total))

        def clean(text: str) -> str:
            if not text:
                return ""
            return ' '.join(text.replace('\n', ' ').replace('\r', ' ').split())

        # =========================
        # 1. CSV Export (clean)
        # =========================
        fieldnames = [
            'ID', 'Priorite', 'Qualification', 'Titre',
            'Client', 'Deadline', 'Budget_Estime', 'Caution',
            'Procedure', 'Domaines_Activite', 'Qualifications',
            'Allotissement', 'Reservation_PME', 'CPS_Source',
            'Description_Technique', 'Description_Fonctionnelle',
            'Requirements', 'URL_Offre', 'Domains', 'Service'
        ]

        with open(csv_path, 'w', encoding='utf-8-sig', newline='') as f:
            writer = csv.DictWriter(
                f, fieldnames=fieldnames, delimiter=';',
                extrasaction='ignore', quoting=csv.QUOTE_ALL
            )
            writer.writeheader()
            for c in relevant:
                qual_parts = []
                if c.matched_domains:
                    qual_parts.append(' / '.join(d.upper() for d in c.matched_domains))
                qual_parts.append(f"Score {c.score_total}")
                qualification = ' - '.join(qual_parts)
                domains_str = ' / '.join(d.upper() for d in c.matched_domains) if c.matched_domains else ''
                service_str = ' / '.join(d.upper() for d in (c.matched_domains or [])[:2]) if c.matched_domains else 'IT'

                writer.writerow({
                    'ID': c.reference,
                    'Priorite': c.priority.value.upper(),
                    'Qualification': qualification,
                    'Titre': clean(c.objet),
                    'Client': c.acheteur or 'Non identifie',
                    'Deadline': c.date_limite or 'Non communiquee',
                    'Budget_Estime': c.estimation_budget or c.budget_estime or '-',
                    'Caution': c.caution_provisoire or '-',
                    'Procedure': c.procedure or '-',
                    'Domaines_Activite': clean(c.domaines_activite) or '-',
                    'Qualifications': clean(c.qualifications) or '-',
                    'Allotissement': clean(c.allotissement) or '-',
                    'Reservation_PME': clean(c.reservation_pme) or '-',
                    'CPS_Source': c.cps_source or '-',
                    'Description_Technique': clean(c.description_technique),
                    'Description_Fonctionnelle': clean(c.description_fonctionnelle),
                    'Requirements': ' | '.join(c.requirements) if c.requirements else '-',
                    'URL_Offre': c.url or '-',
                    'Domains': domains_str or '-',
                    'Service': service_str or 'IT',
                })

        logger.info(f"Export CSV: {len(relevant)} resultats -> {csv_path}")

        # =========================
        # 2. Excel Export (styled)
        # =========================
        wb = Workbook()
        ws = wb.active
        ws.title = "Appels d'offres"

        # Styles
        header_font = Font(name='Calibri', bold=True, size=11, color='FFFFFF')
        header_fill = PatternFill(start_color='003366', end_color='003366', fill_type='solid')
        hot_fill = PatternFill(start_color='C0392B', end_color='C0392B', fill_type='solid')
        warm_fill = PatternFill(start_color='E67E22', end_color='E67E22', fill_type='solid')
        cold_fill = PatternFill(start_color='2980B9', end_color='2980B9', fill_type='solid')
        white_font = Font(name='Calibri', bold=True, size=10, color='FFFFFF')
        wrap_align = Alignment(wrap_text=True, vertical='top')
        center_align = Alignment(horizontal='center', vertical='center')
        top_align = Alignment(vertical='top')
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )

        # Alternating row colors
        even_fill = PatternFill(start_color='F2F7FB', end_color='F2F7FB', fill_type='solid')

        headers = [
            'ID', 'Priorite', 'Qualification', 'Titre',
            'Client', 'Deadline', 'Budget Estime', 'Caution',
            'Procedure', 'Domaines Activite', 'Qualifications',
            'Allotissement', 'Reservation PME', 'CPS Source',
            'Description Technique', 'Description Fonctionnelle',
            'Requirements', 'URL Offre', 'Domains', 'Service'
        ]
        widths = [18, 12, 25, 50, 35, 15, 25, 20, 25, 40, 30, 20, 25, 15, 65, 55, 55, 50, 18, 16]

        # Write headers
        for col, (header, width) in enumerate(zip(headers, widths), 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = thin_border
            ws.column_dimensions[get_column_letter(col)].width = width

        # Write data
        for i, c in enumerate(relevant, 1):
            row = i + 1
            priority_fill = {
                'HOT': hot_fill, 'WARM': warm_fill, 'COLD': cold_fill
            }.get(c.priority.value.upper())

            qual_parts = []
            if c.matched_domains:
                qual_parts.append(' / '.join(d.upper() for d in c.matched_domains))
            qual_parts.append(f"Score {c.score_total}")
            qualification = ' - '.join(qual_parts)
            domains_str = ' / '.join(d.upper() for d in c.matched_domains) if c.matched_domains else '-'
            service_str = ' / '.join(d.upper() for d in (c.matched_domains or [])[:2]) if c.matched_domains else 'IT'

            data = [
                c.reference,
                c.priority.value.upper(),
                qualification,
                clean(c.objet),
                c.acheteur or 'Non identifie',
                c.date_limite or 'Non communiquee',
                c.estimation_budget or c.budget_estime or '-',
                c.caution_provisoire or '-',
                c.procedure or '-',
                clean(c.domaines_activite) or '-',
                clean(c.qualifications) or '-',
                clean(c.allotissement) or '-',
                clean(c.reservation_pme) or '-',
                c.cps_source or '-',
                clean(c.description_technique),
                clean(c.description_fonctionnelle),
                ' | '.join(c.requirements) if c.requirements else '-',
                c.url or '-',
                domains_str,
                service_str,
            ]

            for col, val in enumerate(data, 1):
                cell = ws.cell(row=row, column=col, value=val)
                cell.border = thin_border

                # Alignment
                if col in (15, 16, 17):  # Long text columns
                    cell.alignment = wrap_align
                elif col in (1, 2, 6, 14):  # Center short columns
                    cell.alignment = center_align
                else:
                    cell.alignment = top_align

                # Priority coloring
                if col == 2 and priority_fill:
                    cell.fill = priority_fill
                    cell.font = white_font

                # Alternating rows
                if i % 2 == 0 and col != 2:
                    cell.fill = even_fill

            # Row height based on content
            ws.row_dimensions[row].height = 90

        # Freeze header + auto-filter
        ws.freeze_panes = 'A2'
        ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{len(relevant)+1}"

        wb.save(xlsx_path)
        logger.info(f"Export Excel: {len(relevant)} resultats -> {xlsx_path}")

        return xlsx_path

    # ----- Full pipeline -----
    def run(
        self,
        csv_path: str,
        generate: bool = True,
        min_generate_priority: Priority = Priority.WARM,
        nlp: bool = True,
    ) -> Dict[str, Any]:
        """
        Run the complete pipeline.

        Args:
            csv_path: Path to the scraped CSV (semicolon-delimited, utf-8-sig)
            generate: Whether to generate DOCX dossiers
            min_generate_priority: Minimum priority for dossier generation
            nlp: Whether to run NLP analysis

        Returns:
            Dict with stats and paths
        """
        timestamp = datetime.now().isoformat(sep=' ', timespec='seconds')
        print(f"\n{'='*80}")
        print(f"  PIPELINE ALEXSYS - Analyse des marches publics")
        print(f"  {timestamp}")
        print(f"{'='*80}")

        results = {"timestamp": timestamp, "stages": {}}
        n_stages = 5
        if self.enrich_cps_flag:
            n_stages += 1
        if self.use_rag:
            n_stages += 1
        stage = 0

        # Stage 1: Load
        stage += 1
        print(f"\n[{stage}/{n_stages}] Chargement du CSV...")
        total = self.load_csv(csv_path)
        results["stages"]["load"] = {"total": total, "csv": csv_path}
        print(f"  -> {total} consultations chargees")

        # Stage 2: Score
        stage += 1
        print(f"\n[{stage}/{n_stages}] Scoring IT (keywords ponderes)...")
        score_stats = self.score_all()
        results["stages"]["score"] = score_stats
        print(f"  -> HOT: {score_stats['hot']} | WARM: {score_stats['warm']} | COLD: {score_stats['cold']} | Exclus: {score_stats['excluded']}")

        # Stage 2.5: CPS Enrichment (if enabled)
        if self.enrich_cps_flag:
            stage += 1
            relevant_count = score_stats['hot'] + score_stats['warm'] + score_stats['cold']
            print(f"\n[{stage}/{n_stages}] Enrichissement CPS / pages detail ({relevant_count} consultations)...")
            enrich_stats = self.enrich_from_cps()
            results["stages"]["enrichment"] = enrich_stats
            print(f"  -> {enrich_stats['enriched']} enrichis (PRADO: {enrich_stats['prado']}, BDC: {enrich_stats['bdc']})")
            if enrich_stats['failed'] > 0:
                print(f"  -> {enrich_stats['failed']} echoues")

        # Stage 3: NLP + Descriptions
        stage += 1
        print(f"\n[{stage}/{n_stages}] Analyse NLP + generation descriptions...")
        analyzed = self.analyze_nlp()
        results["stages"]["nlp"] = {"analyzed": analyzed}
        print(f"  -> {analyzed} consultations analysees")

        # Stage 3.5: RAG (if enabled) - overwrites NLP descriptions with LLM-generated
        if self.use_rag:
            stage += 1
            print(f"\n[{stage}/{n_stages}] RAG - Generation intelligente (LLM)...")
            rag_stats = self.analyze_rag()
            results["stages"]["rag"] = rag_stats
            print(f"  -> {rag_stats['rag_enriched']} descriptions generees par LLM")
            if rag_stats['rag_failed'] > 0:
                print(f"  -> {rag_stats['rag_failed']} echoues (fallback NLP)")

        # Stage 4: Generate
        stage += 1
        if generate:
            targets = score_stats['hot']
            if min_generate_priority == Priority.WARM:
                targets += score_stats['warm']
            elif min_generate_priority == Priority.COLD:
                targets += score_stats['warm'] + score_stats['cold']
            print(f"\n[{stage}/{n_stages}] Generation dossiers DOCX ({targets} cibles)...")
            paths = self.generate_dossiers(min_generate_priority)
            results["stages"]["generation"] = {
                "generated": len(paths),
                "paths": paths,
                "output_dir": self.output_dir
            }
            print(f"  -> {len(paths)} dossiers generes dans {self.output_dir}/")
        else:
            print(f"\n[{stage}/{n_stages}] Generation desactivee")
            results["stages"]["generation"] = {"skipped": True}

        # Stage 5/6: Export
        stage += 1
        print(f"\n[{stage}/{n_stages}] Export Excel des resultats...")
        xlsx_out = self.export_results()
        results["stages"]["export"] = {"xlsx": xlsx_out}
        print(f"  -> Export: {xlsx_out}")

        # Summary
        relevant = score_stats['hot'] + score_stats['warm'] + score_stats['cold']
        print(f"\n{'='*80}")
        print(f"  PIPELINE TERMINE")
        print(f"  Consultations IT pertinentes: {relevant}/{total}")
        print(f"  A soumettre en priorite: {score_stats['hot']} HOT + {score_stats['warm']} WARM")
        if generate:
            print(f"  Dossiers generes: {results['stages']['generation']['generated']}")
        print(f"{'='*80}\n")

        results["summary"] = {
            "total": total,
            "relevant": relevant,
            "hot": score_stats['hot'],
            "warm": score_stats['warm'],
            "cold": score_stats['cold'],
        }

        return results


# ============================================================
# CLI entry point
# ============================================================

if __name__ == "__main__":
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s [%(levelname)s] %(message)s',
        datefmt='%H:%M:%S',
    )

    import argparse
    parser = argparse.ArgumentParser(description="Pipeline Alexsys - Analyse marches publics")
    parser.add_argument("csv", nargs='?',
                        default="data/appels_offres_profond_20260310_100313.csv",
                        help="Path to the scraped CSV file")
    parser.add_argument("--no-generate", action="store_true",
                        help="Skip dossier generation")
    parser.add_argument("--no-nlp", action="store_true",
                        help="Skip NLP analysis")
    parser.add_argument("--no-enrich", action="store_true",
                        help="Skip CPS/detail page enrichment")
    parser.add_argument("--rag", action="store_true",
                        help="Enable RAG (LLM-generated descriptions via Ollama/OpenAI)")
    parser.add_argument("--hot-only", action="store_true",
                        help="Generate dossiers only for HOT items")
    parser.add_argument("--all", action="store_true",
                        help="Generate dossiers for all matched items (HOT+WARM+COLD)")
    parser.add_argument("-o", "--output", default="dossiers_generes",
                        help="Output directory for dossiers")
    args = parser.parse_args()

    min_priority = Priority.WARM
    if args.hot_only:
        min_priority = Priority.HOT
    elif args.all:
        min_priority = Priority.COLD

    pipeline = Pipeline(
        output_dir=args.output,
        use_nlp=not args.no_nlp,
        enrich_cps=not args.no_enrich,
        use_rag=args.rag,
    )

    try:
        results = pipeline.run(
            csv_path=args.csv,
            generate=not args.no_generate,
            min_generate_priority=min_priority,
            nlp=not args.no_nlp,
        )
        # Save full results JSON
        keep_ts = os.environ.get("PIPELINE_TIMESTAMP_EXPORTS", "0").strip() == "1"
        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        results_path = f"data/pipeline_report_{ts}.json" if keep_ts else "data/pipeline_report_latest.json"
        with open(results_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, ensure_ascii=False, indent=2, default=str)
        print(f"Rapport complet: {results_path}")
    except KeyboardInterrupt:
        print("\n\nPipeline interrompu par l'utilisateur")
    except Exception as e:
        print(f"\nERREUR: {e}")
        logger.exception("Pipeline error")
        sys.exit(1)
